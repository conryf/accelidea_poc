from contextlib import asynccontextmanager
import hashlib
import io
import os
import re
import zipfile

import anthropic
from dotenv import load_dotenv
import aiosqlite

load_dotenv()
import pytesseract
from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile
from typing import List
from fastapi.responses import HTMLResponse, Response
from pdf2image import convert_from_bytes

DB_PATH = "poc.db"
TEMPLATES_DIR = "templates"
PACKAGE_FILES_DIR = "package_files"

_anthropic = anthropic.AsyncAnthropic()

# (column_name, search label in OCR text)
PACKAGE_FIELDS: list[tuple[str, str]] = [
    ("doc_title",                  "Document Title"),
    ("doc_number",                 "Document Number"),
    ("author_name",                "Author"),
    ("review_date",                "Review Date"),
    ("revision",                   "Revision"),
    ("product_name",               "Product Name"),
    ("product_number",             "Product Number"),
    ("project_name",               "Project Name"),
    ("project_number",             "Project Number"),
    ("description_of_change",      "Description of Change"),
    ("location",                   "Enter location"),
    ("signatory_name",             "Signatory Name"),
    ("scope",                      "Scope"),
    ("test_plan_exceptions",       "Exceptions"),
    ("test_plan_description",      "Test Plan Description"),
    ("test_logistics",             "Test Logistics"),
    ("third_party_testing",        "Third Party Testing"),
    ("company_name",               "Company Name"),
    ("contact_name",               "Contact Name"),
    ("rfp_submission_date",        "Submission Date"),
    ("rfp_overview",               "Overview"),
    ("rfp_company_background",     "Company Background"),
    ("rfp_project_scope",          "Project Scope"),
    ("rfp_project_stages",         "Project Stages"),
    ("rfp_manufacturing_dev",      "Manufacturing Development"),
    ("rfp_process_development",    "Process Development"),
    ("rfp_stability_testing",      "Stability Testing"),
    ("rfp_other_services",         "Other Services"),
    ("rfp_classification",         "Classification"),
    ("rfp_unique_characteristics", "Unique Characteristics"),
]

PACKAGE_FIELDS_DDL = """
    doc_title TEXT,
    doc_number TEXT,
    author_name TEXT,
    review_date DATE,
    revision TEXT,
    product_name TEXT,
    product_number TEXT,
    project_name TEXT,
    project_number TEXT,
    description_of_change TEXT,
    location TEXT,
    signatory_name TEXT,
    scope TEXT,
    test_plan_exceptions TEXT,
    test_plan_description TEXT,
    test_logistics TEXT,
    third_party_testing TEXT,
    company_name TEXT,
    contact_name TEXT,
    rfp_submission_date TEXT,
    rfp_overview TEXT,
    rfp_company_background TEXT,
    rfp_project_scope TEXT,
    rfp_project_stages TEXT,
    rfp_manufacturing_dev BOOLEAN,
    rfp_process_development BOOLEAN,
    rfp_stability_testing BOOLEAN,
    rfp_other_services TEXT,
    rfp_classification TEXT,
    rfp_unique_characteristics TEXT
"""

# (search_text, field_col)
# If search_text matches "Label: [placeholder]", the full string is replaced with "Label: <value>".
# Otherwise the search_text itself is replaced with the value.
TEMPLATE_SUBSTITUTIONS: dict[str, list[tuple[str, str]]] = {
    "Cytotoxicity Test Plan.docx": [
        ("<Add Device Name>", "product_name"),
        ("<company name>",    "company_name"),
        ("<trade name>",      "product_name"),
    ],
    "DVT Test Plan.docx": [
        ("<Description of Change>",                        "description_of_change"),
        ("<Document Number>",                              "doc_number"),
        ("<Document Title>",                               "doc_title"),
        ("<Enter location>",                               "location"),
        ("<Name of Product being Verified, Prod Number>",  "product_name"),
        ("<Project Name and/or Number>",                   "project_name"),
        ("<Revision>",                                     "revision"),
        ("<Signatory Name>",                               "signatory_name"),
        ("<insert product name>",                          "product_name"),
        ("<insert project scope>",                         "scope"),
        ("<product>",                                      "product_name"),
        ("<scope>",                                        "scope"),
        ("<scope of release>",                             "scope"),
        ("<Detail exceptions / limitations applicable to this Test Plan>", "test_plan_exceptions"),
        ("<Outline any third-party testing that is required: test houses, labs, certified facilities, etc.>", "third_party_testing"),
        ("[title]",                                        "doc_title"),
    ],
    "RFP Template.docx": [
        ("[Provide a brief overview of the project and its objectives, emphasizing the importance of compliance with FDA regulations and quality standards.]", "rfp_overview"),
        ("[Provide a brief overview of your company, including product portfolio, product pipeline, and financing details.]",                                   "rfp_company_background"),
        ("[Define the scope of the project, including the type of medical device.]",                                                                           "rfp_project_scope"),
        ("[Specify the stages of the project (design, development, manufacturing, testing, etc.).]",                                                           "rfp_project_stages"),
        ("Manufacturing Development: [Yes/No]",    "rfp_manufacturing_dev"),
        ("Process Development: [Yes/No]",          "rfp_process_development"),
        ("Stability Testing: [Yes/No]",            "rfp_stability_testing"),
        ("[Specify other services required]",      "rfp_other_services"),
        ("[Select appropriate classification: API, Biologically Produced Macromolecule, Non-aqueous Solvents, etc.]", "rfp_classification"),
    ],
}


def _build_replacements(template_name: str, extracted: dict) -> list[tuple[str, str]]:
    """Return (find, replace) pairs for a given template, skipping fields with no value."""
    result = []
    for find, field_col in TEMPLATE_SUBSTITUTIONS.get(template_name, []):
        val = extracted.get(field_col)
        if not val:
            continue
        # If the find text is "Label: [placeholder]", keep the label prefix
        m = re.match(r'^(.+?:\s*)\[', find)
        if m:
            result.append((find, m.group(1) + val))
        else:
            result.append((find, val))
    return result


def _replace_in_paragraph(para, find: str, replace: str) -> None:
    full = "".join(r.text for r in para.runs)
    if find not in full:
        return
    # Try single-run replacement first to preserve per-run formatting
    for run in para.runs:
        if find in run.text:
            run.text = run.text.replace(find, replace)
            return
    # Placeholder spans multiple runs — consolidate into first run
    para.runs[0].text = full.replace(find, replace)
    for run in para.runs[1:]:
        run.text = ""


def _apply_substitutions(doc: Document, replacements: list[tuple[str, str]]) -> None:
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for para in all_paragraphs:
        for find, replace in replacements:
            _replace_in_paragraph(para, find, replace)


def populate_templates(extracted: dict, package_id: int) -> list[str]:
    """Fill all three templates with extracted values, save to package_files/, return saved paths."""
    os.makedirs(PACKAGE_FILES_DIR, exist_ok=True)
    saved = []
    for template_name in TEMPLATE_SUBSTITUTIONS:
        template_path = os.path.join(TEMPLATES_DIR, template_name)
        doc = Document(template_path)
        replacements = _build_replacements(template_name, extracted)
        _apply_substitutions(doc, replacements)

        stem = os.path.splitext(template_name)[0]
        out_name = f"{stem}_{package_id}.docx"
        out_path = os.path.join(PACKAGE_FILES_DIR, out_name)
        doc.save(out_path)
        saved.append(out_path)
    return saved


async def extract_field_with_source(
    per_doc_texts: dict[str, str], label: str
) -> tuple[str | None, str | None]:
    """Try each document in turn. Returns (value, source_filename) or (None, None)."""
    pattern = re.compile(
        rf'{re.escape(label)}\s*[:\-]?\s*(.+?)(?=\n|$)',
        re.IGNORECASE,
    )

    # Regex pass — try each document individually
    for filename, text in per_doc_texts.items():
        match = pattern.search(text)
        if match:
            value = match.group(1).strip()
            if value:
                return value, filename

    # LLM fallback — try each document individually
    for filename, text in per_doc_texts.items():
        message = await _anthropic.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=256,
            system=(
                "You are a document field extractor. "
                "Given OCR text from a document, extract the value for the requested field. "
                "Reply with only the extracted value, or the single word NULL if the field is not present."
            ),
            messages=[
                {
                    "role": "user",
                    "content": f"Field to extract: {label}\n\nDocument text:\n{text[:4000]}",
                }
            ],
        )
        value = message.content[0].text.strip()
        if value.upper() != "NULL":
            return value, filename

    return None, None


@asynccontextmanager
async def lifespan(app: FastAPI):
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute("""
            CREATE TABLE IF NOT EXISTS uploads(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT NOT NULL,
                page INTEGER NOT NULL,
                text TEXT NOT NULL,
                package_id INTEGER NOT NULL,
                md5 TEXT NOT NULL
            )
        """)
        await db.execute(f"""
            CREATE TABLE IF NOT EXISTS packages(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                {PACKAGE_FIELDS_DDL}
            )
        """)
        await db.execute("""
            CREATE TABLE IF NOT EXISTS package_fields(
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                package_id INTEGER NOT NULL,
                field_name TEXT NOT NULL,
                source TEXT NOT NULL
            )
        """)
        # add any missing columns to existing databases
        for col, _ in PACKAGE_FIELDS:
            try:
                await db.execute(f"ALTER TABLE packages ADD COLUMN {col} TEXT")
            except Exception:
                pass
        await db.commit()
    yield


app = FastAPI(lifespan=lifespan)


@app.get("/")
def read_root():
    return {"Mmmmmmm": "Yelllllo"}


@app.get("/items/{item_id}")
def read_item(item_id: int, q: str | None = None):
    return {"item_id": item_id, "q": q}


@app.get("/upload", response_class=HTMLResponse)
async def upload_form():
    return """
<!DOCTYPE html>
<html>
<head><title>Upload PDF</title></head>
<body>
  <h2>Upload PDFs</h2>
  <form action="/upload" method="post" enctype="multipart/form-data">
    <input type="file" name="files" accept="application/pdf" multiple required>
    <button type="submit">Upload</button>
  </form>
</body>
</html>
"""


@app.post("/upload")
async def upload_file(files: List[UploadFile] = File(...)):
    # Validate all files up front
    for file in files:
        if file.content_type != "application/pdf":
            raise HTTPException(status_code=400, detail=f"'{file.filename}' is not a PDF")

    # Read contents and check for duplicates before doing any heavy work
    file_data: list[tuple[str, bytes, str]] = []  # (filename, contents, md5)
    for file in files:
        contents = await file.read()
        md5 = hashlib.md5(contents).hexdigest()
        file_data.append((file.filename, contents, md5))

    async with aiosqlite.connect(DB_PATH) as db:
        for filename, _, md5 in file_data:
            async with db.execute("SELECT filename FROM uploads WHERE md5 = ? LIMIT 1", (md5,)) as cursor:
                existing = await cursor.fetchone()
            if existing:
                raise HTTPException(status_code=409, detail=f"'{filename}' was already uploaded as '{existing[0]}'")

        # OCR all files, accumulating pages and per-document text
        per_doc_texts: dict[str, str] = {}
        upload_rows: list[tuple] = []
        for filename, contents, md5 in file_data:
            images = convert_from_bytes(contents)
            doc_pages = []
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                doc_pages.append(text)
                upload_rows.append((filename, i + 1, text, md5))
            per_doc_texts[filename] = "\n".join(doc_pages)

        extracted: dict[str, str | None] = {}
        field_sources: list[tuple[str, str]] = []  # (field_name, source_filename)
        for col, label in PACKAGE_FIELDS:
            value, source = await extract_field_with_source(per_doc_texts, label)
            extracted[col] = value
            if value is not None and source is not None:
                field_sources.append((label, source))

        package_name = ", ".join(f for f, _, _ in file_data)
        field_cols = ", ".join(["name"] + [col for col, _ in PACKAGE_FIELDS])
        placeholders = ", ".join(["?"] * (1 + len(PACKAGE_FIELDS)))
        field_values = [package_name] + [extracted[col] for col, _ in PACKAGE_FIELDS]

        cursor = await db.execute(
            f"INSERT INTO packages ({field_cols}) VALUES ({placeholders})",
            field_values,
        )
        package_id = cursor.lastrowid

        await db.executemany(
            "INSERT INTO uploads (filename, page, text, package_id, md5) VALUES (?, ?, ?, ?, ?)",
            [(filename, page, text, package_id, md5) for filename, page, text, md5 in upload_rows],
        )
        await db.executemany(
            "INSERT INTO package_fields (package_id, field_name, source) VALUES (?, ?, ?)",
            [(package_id, field_name, source) for field_name, source in field_sources],
        )
        await db.commit()

    os.makedirs("artifacts", exist_ok=True)
    for filename, contents, _ in file_data:
        artifact_path = os.path.join("artifacts", f"package{package_id}_{filename}")
        with open(artifact_path, "wb") as f:
            f.write(contents)

    saved_paths = populate_templates(extracted, package_id)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in saved_paths:
            zf.write(path, os.path.basename(path))
    zip_buffer.seek(0)

    return Response(
        content=zip_buffer.read(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="package_{package_id}.zip"'},
    )
